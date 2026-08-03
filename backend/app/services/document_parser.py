"""文档解析服务：从 Word/PDF/TXT/Markdown 中提取纯文本。"""
from pathlib import Path

from docx import Document
import pdfplumber


def extract_text(file_path: str, file_type: str) -> str:
    """按文件类型提取纯文本，失败抛出 ValueError。"""
    if file_type == "docx":
        return _extract_docx(file_path)
    if file_type == "pdf":
        return _extract_pdf(file_path)
    if file_type in ("txt", "md"):
        return _extract_plain(file_path)
    raise ValueError(f"不支持的文件类型: {file_type}")


def _extract_docx(file_path: str) -> str:
    """提取 Word 文档文本（段落 + 表格）。"""
    doc = Document(file_path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("Word 文档未提取到文本内容")
    return text


def _extract_pdf(file_path: str) -> str:
    """提取 PDF 文本（逐页）。"""
    with pdfplumber.open(file_path) as pdf:
        parts = [page.extract_text() or "" for page in pdf.pages]
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("PDF 未提取到文本内容（扫描件暂不支持 OCR）")
    return text


def _extract_plain(file_path: str) -> str:
    """提取 TXT/Markdown 文本（优先 UTF-8，兼容 GB18030）。"""
    data = Path(file_path).read_bytes()
    for encoding in ("utf-8", "gb18030"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise ValueError("文本编码无法识别，请使用 UTF-8 或 GBK 编码")
    if not text.strip():
        raise ValueError("文件内容为空")
    return text
